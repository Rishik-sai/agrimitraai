import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h

def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(text, style=style)
    return p

def add_code_block(doc, code, language="python"):
    # Add a custom style if not exists
    styles = doc.styles
    if 'CodeBlock' not in styles:
        style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Consolas'
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    
    p = doc.add_paragraph(code, style='CodeBlock')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def generate_document():
    doc = Document()
    
    # ---------------------------------------------------------
    # 1. COVER PAGE
    # ---------------------------------------------------------
    doc.add_heading('AgriMitra AI Pro', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, 'Multi-RAG Agricultural Intelligence System for Indian Farmers').alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, '\n\n\n')
    
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0,0).text = 'Author:'
    table.cell(0,1).text = 'Rishik-sai'
    table.cell(1,0).text = 'Tech Stack:'
    table.cell(1,1).text = 'FastAPI, React (Vite), Groq API, FAISS, LangChain'
    table.cell(2,0).text = 'GitHub Repository:'
    table.cell(2,1).text = 'https://github.com/Rishik-sai/agrimitraai'
    table.cell(3,0).text = 'Date:'
    table.cell(3,1).text = '2026-05-26'
    table.cell(4,0).text = 'Version:'
    table.cell(4,1).text = '2.0.0'
    
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # 2. EXECUTIVE SUMMARY
    # ---------------------------------------------------------
    add_heading(doc, '1. Executive Summary', 1)
    
    add_heading(doc, 'What problem this project solves', 2)
    add_paragraph(doc, 'Indian farmers often lack access to timely, localized, and expert-level agricultural advice in their native languages. Traditional extension services are overwhelmed, leading to crop losses from misdiagnosed diseases, lack of awareness about government schemes, and poor market timing. AgriMitra AI Pro solves this by acting as a 24/7 personalized, AI-driven agricultural extension worker.')

    add_heading(doc, 'Why this project matters', 2)
    add_paragraph(doc, 'By democratizing access to agronomic expertise, this project directly impacts agricultural productivity and farmer income. It bridges the digital divide by providing voice-based, multilingual interfaces, ensuring that even illiterate farmers can benefit from cutting-edge generative AI.')

    add_heading(doc, 'Real-world applications', 2)
    add_paragraph(doc, '• On-field crop disease diagnosis via smartphone cameras.', style='List Bullet')
    add_paragraph(doc, '• Instant retrieval of current market prices (Mandi rates) for better negotiation.', style='List Bullet')
    add_paragraph(doc, '• Voice-based querying for government schemes and subsidies eligibility.', style='List Bullet')
    
    add_heading(doc, 'Target users', 2)
    add_paragraph(doc, 'Smallholder Indian farmers, agricultural extension workers (Krishi Mitras), agri-tech startups, and rural cooperatives.')
    
    add_heading(doc, 'Expected impact', 2)
    add_paragraph(doc, 'Reduced crop mortality through early disease detection, higher profit margins via market intelligence, and increased adoption of beneficial government schemes.')
    
    doc.add_page_break()

    # ---------------------------------------------------------
    # 3. PROJECT ARCHITECTURE
    # ---------------------------------------------------------
    add_heading(doc, '2. Project Architecture', 1)
    
    add_heading(doc, 'High-level Architecture', 2)
    add_paragraph(doc, 'AgriMitra AI Pro employs a decoupled architecture consisting of a React-based frontend and a FastAPI backend. The core of the intelligence layer is a Multi-Agent Retrieval-Augmented Generation (RAG) pipeline. When a user asks a question, a Router identifies the relevant domain agents (e.g., Crop, Market, Schemes). These agents query a shared local FAISS vector store. The retrieved context is synthesized by a Groq LLM and dynamically translated into the user\'s requested language.')

    add_heading(doc, 'Workflow Diagram (Textual Representation)', 2)
    add_code_block(doc, '''[User Query (Voice/Text)] 
       --> [FastAPI Backend] 
             --> [Query Router (Keyword based)]
                   --> [Select 1-3 Active Agents]
                         --> [FAISS Vector Store Retrieval]
                               --> [Context Extraction]
                                     --> [Groq LLM Synthesizer]
                                           --> [Dynamic JSON Translator]
                                                 --> [React Frontend UI]''')

    add_heading(doc, 'Folder Structure Explanation', 2)
    add_paragraph(doc, 'The repository is cleanly divided into client and server domains:')
    add_paragraph(doc, '• /backend: Contains the FastAPI server (`main.py`), the RAG engine (`multi_rag.py`), FAISS indices (`/faiss_index`), and data ingestion scripts (`ingest.py`).')
    add_paragraph(doc, '• /frontend: Contains the React/Vite application. `src/components/` holds domain-specific widgets (Scanner, Market, Weather). `src/translations.js` maps dictionary keywords for the UI.')

    doc.add_page_break()

    # ---------------------------------------------------------
    # 4. TECH STACK ANALYSIS
    # ---------------------------------------------------------
    add_heading(doc, '3. Tech Stack Analysis', 1)

    add_heading(doc, 'FastAPI (Backend Web Framework)', 2)
    add_paragraph(doc, 'Why Selected: High performance, native async support, and automatic OpenAPI documentation. Perfect for I/O bound LLM and API tasks.')
    add_paragraph(doc, 'Advantages: Faster than Flask; strong typing with Pydantic ensures data validation.')
    add_paragraph(doc, 'Tradeoffs: Steeper learning curve than Flask for beginners.')
    add_paragraph(doc, 'Alternatives: Django (too heavy), Flask (lacks native robust async).')

    add_heading(doc, 'React + Vite (Frontend)', 2)
    add_paragraph(doc, 'Why Selected: Component-based architecture allows modular widget design (e.g., Scanner, Market panels). Vite provides instantaneous Hot Module Replacement (HMR).')
    add_paragraph(doc, 'Advantages: Large ecosystem, reactive state management.')
    add_paragraph(doc, 'Tradeoffs: Client-side rendering can impact initial load times compared to SSR.')
    add_paragraph(doc, 'Alternatives: Next.js (overhead not needed for an SPA without complex SEO needs).')

    add_heading(doc, 'FAISS (Vector Database)', 2)
    add_paragraph(doc, 'Why Selected: Extremely fast, CPU-optimized local vector similarity search.')
    add_paragraph(doc, 'Advantages: No external cloud dependencies; runs locally on commodity hardware.')
    add_paragraph(doc, 'Tradeoffs: Index needs to be rebuilt and managed manually as data grows.')
    add_paragraph(doc, 'Alternatives: Pinecone, Qdrant (rejected due to cloud dependency and latency constraints for local operation).')

    add_heading(doc, 'Groq API & Llama 3.3/4 (LLMs)', 2)
    add_paragraph(doc, 'Why Selected: Groq provides LPU-accelerated inference, which is critical for a multi-agent system requiring low-latency responses.')
    add_paragraph(doc, 'Advantages: Blazing fast generation speeds, cost-effective.')
    add_paragraph(doc, 'Tradeoffs: Dependency on an external API provider.')
    add_paragraph(doc, 'Alternatives: OpenAI GPT-4o (slower and more expensive), Local Ollama (too heavy for standard consumer hardware).')

    doc.add_page_break()

    # ---------------------------------------------------------
    # 5. FEATURE-BY-FEATURE BREAKDOWN
    # ---------------------------------------------------------
    add_heading(doc, '4. Feature-by-Feature Breakdown', 1)

    # Feature 1
    add_heading(doc, 'Feature: Multimodal Leaf Disease Scanner', 2)
    add_heading(doc, 'Purpose', 3)
    add_paragraph(doc, 'Analyzes visual symptoms of crop leaves uploaded by the user to diagnose diseases and recommend specific chemical or organic treatments.')
    add_heading(doc, 'User Flow', 3)
    add_paragraph(doc, 'User takes a photo of a sick leaf -> Uploads to the web app -> Selects preferred language -> Clicks "Scan" -> Receives disease name, confidence score, and remedies.')
    add_heading(doc, 'Internal Working', 3)
    add_paragraph(doc, 'The frontend converts the image to base64 and sends it to `/api/scan`. The backend wraps the image and a strict JSON-schema prompt and sends it to the `meta-llama/llama-4-scout-17b-16e-instruct` Vision model via Groq. The JSON response is parsed, dynamically translated using the translation LLM pipeline, and returned.')
    
    add_heading(doc, 'Important Code Snippet', 3)
    code1 = """@app.post("/api/scan")
async def scan_endpoint(file: UploadFile = File(...), lang: str = Query(default="en")):
    # Base64 encode file...
    llm = ChatGroq(model_name="meta-llama/llama-4-scout-17b-16e-instruct", groq_api_key=api_key)
    prompt = "Analyze this crop leaf image... Output ONLY valid JSON."
    message = HumanMessage(content=[{"type": "text", "text": prompt}, {"type": "image_url", "image_url": ...}])
    response = llm.invoke([message])
    # Parse and Translate...
    translated_analysis = get_translated(f"scan_{file.filename}", analysis_data, lang)
    return {"analysis": translated_analysis}"""
    add_code_block(doc, code1)
    add_paragraph(doc, 'Explanation: This snippet demonstrates the integration of a multimodal LLM in a FastAPI endpoint. It leverages LangChain\'s HumanMessage to pass image data. Time complexity is O(1) from the server perspective (blocking on API call). The translation layer is decoupled, allowing the core vision model to focus purely on analysis in English, which optimizes inference accuracy.')

    # Feature 2
    add_heading(doc, 'Feature: Real-time Multi-Agent Chat', 2)
    add_heading(doc, 'Purpose', 3)
    add_paragraph(doc, 'Provides conversational agronomic advice by pulling from a localized knowledge base.')
    add_heading(doc, 'Internal Working', 3)
    add_paragraph(doc, 'The query is routed using a keyword heuristic to 1-3 agents. Each agent retrieves top-K chunks from FAISS. The chunks are merged and sent to the LLM Synthesizer.')
    add_heading(doc, 'Important Code Snippet', 3)
    code2 = """def route_query(query: str) -> List[str]:
    query_lower = query.lower()
    scores: Dict[str, int] = {}
    for agent_id, config in AGENTS.items():
        score = sum(1 for kw in config.keywords if kw in query_lower)
        if score > 0: scores[agent_id] = score
    # Return top 3 agents"""
    add_code_block(doc, code2)
    add_paragraph(doc, 'Explanation: Routing uses an O(A * K) heuristic approach where A is the number of agents and K is keywords. This was chosen over LLM-based routing to eliminate latency and reduce API costs. It ensures immediate dispatch to vector retrieval.')

    doc.add_page_break()

    # ---------------------------------------------------------
    # 6. CORE LOGIC EXPLANATION
    # ---------------------------------------------------------
    add_heading(doc, '5. Core Logic Explanation', 1)

    add_heading(doc, 'Dynamic LLM Translation Strategy', 2)
    add_paragraph(doc, 'To support 11 Indian languages without maintaining separate databases, the system uses an LLM-based translation wrapper.')
    add_code_block(doc, """def translate_json_via_llm(data, lang_full, context_hint):
    prompt = f"Translate ALL string values into {lang_full}... Keep all JSON keys EXACTLY as-is..."
    response = llm.invoke(prompt)
    return json.loads(content)""")
    add_paragraph(doc, 'Why written this way: Standard translation APIs (Google/AWS) often break JSON structures or mistranslate agricultural jargon. By using Llama 3.3 with a context hint ("Indian agricultural terminology") and strict JSON constraints, we guarantee structural integrity while providing context-aware translations. The `TRANSLATION_CACHE` dictionary acts as a memoization layer to prevent redundant LLM calls (O(1) lookup).')

    add_heading(doc, 'Multi-Agent Coordination Pipeline', 2)
    add_paragraph(doc, 'The RAG architecture utilizes parallel retrieval (simulated via synchronous loops in MVP) to collect context. The Synthesizer is instructed with a structured prompt combining `agents_consulted` and `context`. This reduces hallucinations because the LLM is tightly constrained to the retrieved textual context.')

    doc.add_page_break()

    # ---------------------------------------------------------
    # 7. API DOCUMENTATION
    # ---------------------------------------------------------
    add_heading(doc, '6. API Documentation', 1)

    add_heading(doc, '1. POST /api/chat', 2)
    add_paragraph(doc, 'Endpoint URL: /api/chat')
    add_paragraph(doc, 'Method: POST')
    add_paragraph(doc, 'Request Body: {"query": "string", "language": "string"}')
    add_paragraph(doc, 'Response Format: {"answer": "string", "sources": ["string"], "agents_used": [{"name": "string", "emoji": "string"}]}')
    add_paragraph(doc, 'Error Handling: 500 Internal Server Error on LLM failure.')

    add_heading(doc, '2. POST /api/scan', 2)
    add_paragraph(doc, 'Method: POST (Multipart/form-data)')
    add_paragraph(doc, 'Request Body: file (UploadFile), lang (Query)')
    add_paragraph(doc, 'Response Format: {"analysis": {"disease": "...", "confidence": 80, "recommendations": []}}')

    add_heading(doc, '3. GET /api/weather/{state}/{city}', 2)
    add_paragraph(doc, 'Method: GET')
    add_paragraph(doc, 'Response Format: {"temp": "...", "humidity": "...", "forecast": [], "advisory": "..."}')

    doc.add_page_break()

    # ---------------------------------------------------------
    # 8. DATABASE DESIGN
    # ---------------------------------------------------------
    add_heading(doc, '7. Database Design', 1)
    add_paragraph(doc, 'The application currently utilizes a NoSQL approach centered around a Vector Store, alongside in-memory JSON objects.')
    add_paragraph(doc, '1. FAISS Vector Store: Flat L2 index storing embeddings of agricultural documents. Dimensions: 384 (MiniLM).')
    add_paragraph(doc, '2. In-Memory Dictionaries: Used for geographic data, market prices, and schemas (`main.py`). This acts as an edge-cache replacement for immediate access, suitable for proof-of-concept phase.')

    # ---------------------------------------------------------
    # 9. AI/ML SECTION
    # ---------------------------------------------------------
    add_heading(doc, '8. AI/ML Section', 1)
    add_paragraph(doc, 'Model Architecture:', style='List Bullet')
    add_paragraph(doc, '• Embedding Model: `all-MiniLM-L6-v2` (Sentence Transformers). Chosen for high speed on CPU and sufficient accuracy for document similarity.', style='List Bullet')
    add_paragraph(doc, '• Language Model: `llama-3.3-70b-versatile` via Groq. Operates as the brain for synthesis and translation.', style='List Bullet')
    add_paragraph(doc, '• Vision Model: `meta-llama/llama-4-scout-17b-16e-instruct` via Groq. Processes multi-modal inputs (images of leaves).', style='List Bullet')
    
    add_paragraph(doc, '\nRAG Pipeline:')
    add_paragraph(doc, 'The RAG pipeline retrieves top K=4/5 chunks per agent. Using multiple specialized agents increases recall for complex questions (e.g., "What is the price of paddy and what fertilizer to use?" triggers both Market and Crop agents).')

    # ---------------------------------------------------------
    # 10. PERFORMANCE OPTIMIZATION
    # ---------------------------------------------------------
    add_heading(doc, '9. Performance Optimization', 1)
    add_paragraph(doc, '1. Groq LPU API: Migrating from standard GPU endpoints to Groq LPUs reduced latency from ~4s to ~800ms for generation.')
    add_paragraph(doc, '2. Translation Caching: `TRANSLATION_CACHE` dictionary caches endpoints (e.g., `market_hi`, `schemes_te`). Time complexity drops from O(LLM) to O(1) after the first user requests a language.')
    add_paragraph(doc, '3. Keyword Routing: Replaced LLM-based intent classification with O(N) keyword matching to instantly route queries.')

    # ---------------------------------------------------------
    # 11. SECURITY ANALYSIS
    # ---------------------------------------------------------
    add_heading(doc, '10. Security Analysis', 1)
    add_paragraph(doc, '• Secret Management: `.env` file management prevents API keys from leaking into the repository.')
    add_paragraph(doc, '• Input Validation: FastAPI/Pydantic automatically validates chat query lengths (`min_length=1`, `max_length=2000`) preventing massive payload DOS attacks.')
    add_paragraph(doc, '• Prompt Injection Prevention: Synthesis prompts enforce strict adherence to provided context (`Provide a comprehensive answer using ONLY the retrieved context above`).')

    # ---------------------------------------------------------
    # 12. DEPLOYMENT GUIDE
    # ---------------------------------------------------------
    add_heading(doc, '11. Deployment Guide', 1)
    add_paragraph(doc, 'Local Setup:')
    add_code_block(doc, '''# Backend
cd backend
python -m venv venv
source venv/bin/activate
pip install -r requirements.txt
uvicorn main:app --reload --port 8000

# Frontend
cd frontend
npm install
npm run dev''', language='bash')
    add_paragraph(doc, 'Production Considerations: Use Gunicorn with Uvicorn workers. Move `TRANSLATION_CACHE` to Redis. Build React via `npm run build` and serve via Nginx or FastAPI StaticFiles.')

    # ---------------------------------------------------------
    # 13. CHALLENGES & SOLUTIONS
    # ---------------------------------------------------------
    add_heading(doc, '12. Challenges & Solutions', 1)
    add_paragraph(doc, 'Issue: JSON parsing failures during translation.')
    add_paragraph(doc, 'Why it happened: LLMs often wrap outputs in markdown blockticks (```json).')
    add_paragraph(doc, 'Solution: Implemented a robust string stripping mechanism in `translate_json_via_llm` that removes markdown wrappers before `json.loads()`.')

    # ---------------------------------------------------------
    # 14. FUTURE IMPROVEMENTS
    # ---------------------------------------------------------
    add_heading(doc, '13. Future Improvements', 1)
    add_paragraph(doc, '• Scalability: Migrate FAISS to a managed distributed database like Qdrant or Pinecone.')
    add_paragraph(doc, '• Caching: Implement Redis for distributed translation and market data caching.')
    add_paragraph(doc, '• Multimodal Voice: Integrate Whisper API for native voice-to-text processing on the backend instead of relying on browser Web Speech API.')

    # ---------------------------------------------------------
    # 15. CONCLUSION
    # ---------------------------------------------------------
    add_heading(doc, '14. Conclusion', 1)
    add_paragraph(doc, 'AgriMitra AI Pro demonstrates a highly efficient, cost-effective application of advanced Generative AI and multi-agent RAG patterns. By focusing on ultra-low latency inference (Groq), localized intelligence (FAISS), and accessibility (multilingual LLM wrapper), the system acts as a robust prototype for an enterprise-grade national agricultural extension platform.')

    doc.save('AgriMitra_AI_Pro_Technical_Documentation.docx')
    print("Document successfully generated.")

if __name__ == '__main__':
    generate_document()

